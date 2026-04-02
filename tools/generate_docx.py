#!/usr/bin/env python3
"""
DeepBrief AI — .docx 報告生成器（精修版）
"""
import json, sys, os, argparse
from datetime import datetime
from pathlib import Path

import re as _re

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENT

# ── 品牌色彩 ──────────────────────────────────────
DEEP_NAVY  = RGBColor(0x00, 0x1A, 0x4E)
NEXI_BLUE  = RGBColor(0x00, 0x30, 0x87)
LIGHT_BLUE = RGBColor(0x00, 0xA3, 0xE0)
TEAL       = RGBColor(0x00, 0xC2, 0xCB)
CHARCOAL   = RGBColor(0x1A, 0x1A, 0x2E)
AMBER      = RGBColor(0xFF, 0xB4, 0x00)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GREY = RGBColor(0xBB, 0xBB, 0xCC)
DATA_BG    = RGBColor(0xE8, 0xF4, 0xFF)


# ── XML helpers ───────────────────────────────────

def set_para_shading(para, color: RGBColor):
    """段落背景色（用於全版色塊）"""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    hex_c = f'{color[0]:02X}{color[1]:02X}{color[2]:02X}'
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_c)
    pPr.append(shd)


def set_cell_bg(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_c = f'{color[0]:02X}{color[1]:02X}{color[2]:02X}'
    shd.set(qn('w:fill'), hex_c)
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def remove_table_borders(table):
    tbl = table._tbl
    # python-docx 1.2+ 不再有 get_or_add_tblPr，改用 lxml 直接操作
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'none')
        tblBorders.append(el)
    tblPr.append(tblBorders)


def set_cell_margins(cell, top=0, bottom=0, left=200, right=200):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)


def set_row_height(row, twips, exact=True):
    trPr = row._tr.get_or_add_trPr()
    trH = OxmlElement('w:trHeight')
    trH.set(qn('w:val'), str(twips))
    if exact:
        trH.set(qn('w:hRule'), 'exact')
    trPr.append(trH)


def add_run(para, text, bold=False, size=None, color=None, italic=False):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


# ── Markdown → docx helpers ──────────────────────

def _strip_md_block(text: str) -> str:
    """
    將 Claude 偶爾輸出的 markdown 語法轉換/清理為純文字，
    同時保留段落結構（雙換行）。
    處理：### 標題、**粗體**、*斜體*、--- 分隔線、markdown 表格、- 項目符號
    """
    lines = text.split('\n')
    cleaned = []
    in_table = False

    for line in lines:
        stripped = line.strip()

        # 跳過 markdown 分隔線 (---, ***, ___)
        if _re.match(r'^[-*_]{3,}\s*$', stripped):
            continue

        # 跳過 markdown 表格分隔列 (|---|---|)
        if _re.match(r'^\|[\s\-:|]+\|$', stripped):
            in_table = True
            continue

        # markdown 表格資料列 → 提取為文字
        if stripped.startswith('|') and stripped.endswith('|'):
            cells = [c.strip() for c in stripped.strip('|').split('|')]
            cells = [c for c in cells if c]
            if cells:
                cleaned.append('　'.join(cells))
            in_table = True
            continue

        in_table = False

        # ### 標題 → 移除 # 前綴，保留文字
        if stripped.startswith('#'):
            stripped = _re.sub(r'^#{1,6}\s*', '', stripped)

        # **粗體** 和 *斜體* → 移除 markdown 標記
        stripped = _re.sub(r'\*\*(.+?)\*\*', r'\1', stripped)
        stripped = _re.sub(r'\*(.+?)\*', r'\1', stripped)
        stripped = _re.sub(r'__(.+?)__', r'\1', stripped)
        stripped = _re.sub(r'_(.+?)_', r'\1', stripped)

        # `code` → 移除反引號
        stripped = _re.sub(r'`(.+?)`', r'\1', stripped)

        # - 項目符號 → 替換為中文符號
        if _re.match(r'^[-•]\s+', stripped):
            stripped = _re.sub(r'^[-•]\s+', '▸ ', stripped)
        elif _re.match(r'^\d+\.\s+', stripped):
            pass  # 數字列表保留原樣

        cleaned.append(stripped)

    return '\n'.join(cleaned)


def _render_rich_paragraph(doc, text, indent=True):
    """
    渲染一個段落，支援 **粗體** 內嵌語法。
    其餘 markdown 語法應已被 _strip_md_block 清除。
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    if indent:
        p.paragraph_format.first_line_indent = Pt(22)

    # 切分 **粗體** 片段
    parts = _re.split(r'(\*\*.+?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            add_run(p, part[2:-2], bold=True, size=11, color=CHARCOAL)
        else:
            add_run(p, part, size=11, color=CHARCOAL)
    return p


# ── Table of Contents ────────────────────────────

def build_toc(doc, sections, summary_exists=True):
    """生成靜態目錄頁"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(20)
    add_run(p, '目　錄', bold=True, size=22, color=NEXI_BLUE)

    items = []
    if summary_exists:
        items.append(('主管摘要', 0))
    for sec in sections:
        items.append((sec.get('title', ''), 0))
    items.append(('研究限制與風險提示', 0))
    items.append(('附錄：參考資料來源', 0))

    for title, _ in items:
        tp = doc.add_paragraph()
        tp.paragraph_format.space_after = Pt(6)
        tp.paragraph_format.left_indent = Cm(0.5)
        add_run(tp, '▸  ', color=TEAL, size=11)
        add_run(tp, title, size=11, color=CHARCOAL)

    # 裝飾底線
    lp = doc.add_paragraph()
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lp.paragraph_format.space_before = Pt(16)
    lr = lp.add_run('─' * 40)
    lr.font.size = Pt(9)
    lr.font.color.rgb = RGBColor(0xCC, 0xCC, 0xDD)

    doc.add_page_break()


# ── Data Table ───────────────────────────────────

def build_data_table(doc, table_data):
    """渲染 JSON 表格為 Word 表格（深藍表頭、斑馬紋列）"""
    headers = table_data.get('headers', [])
    rows = table_data.get('rows', [])
    if not headers or not rows:
        return

    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.autofit = True

    # 表頭
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_bg(cell, NEXI_BLUE)
        set_cell_margins(cell, top=60, bottom=60, left=120, right=120)
        cell.paragraphs[0].clear()
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        add_run(p, str(h), bold=True, size=9, color=WHITE)

    # 資料列
    zebra = RGBColor(0xF5, 0xF8, 0xFC)
    for r_idx, row_data in enumerate(rows):
        for c_idx in range(n_cols):
            cell = table.rows[r_idx + 1].cells[c_idx]
            if r_idx % 2 == 1:
                set_cell_bg(cell, zebra)
            set_cell_margins(cell, top=40, bottom=40, left=120, right=120)
            cell.paragraphs[0].clear()
            val = row_data[c_idx] if c_idx < len(row_data) else ''
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            add_run(p, str(val), size=9, color=CHARCOAL)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ── Cover Page ────────────────────────────────────

def build_cover_page(doc, meta):
    """精修封面：使用多列表格實現全版深色設計"""
    PAGE_W_CM = 21.0
    PAGE_H_CM = 29.7

    # 計算可用寬度（頁面寬 - 左右邊距）
    section = doc.sections[0]
    left_margin_cm  = section.left_margin / 914400 * 2.54
    right_margin_cm = section.right_margin / 914400 * 2.54
    usable_w_cm = PAGE_W_CM - left_margin_cm - right_margin_cm

    # 全版封面表格（1行1列）
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    remove_table_borders(table)

    col_w = int(usable_w_cm * 360000)  # EMU
    table.columns[0].width = col_w

    row = table.rows[0]
    set_row_height(row, 11800, exact=True)   # ~20.8cm 高

    cell = table.cell(0, 0)
    set_cell_bg(cell, DEEP_NAVY)
    set_cell_margins(cell, top=600, bottom=400, left=500, right=500)

    # 清除預設段落
    cell.paragraphs[0].clear()

    def cover_para(text, size, bold=False, color=WHITE,
                   align=WD_ALIGN_PARAGRAPH.CENTER,
                   space_before=0, space_after=6):
        p = cell.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        p.paragraph_format.line_spacing = Pt(size * 1.3)
        add_run(p, text, bold=bold, size=size, color=color)
        return p

    # 品牌標誌列
    cover_para('Dolphin.Ai', 13, bold=True, color=TEAL,
               space_before=28, space_after=2)
    cover_para('Dive in Your Market', 9, color=LIGHT_BLUE, space_after=24)

    # 裝飾線（以空格＋底線色模擬）
    line_p = cell.add_paragraph()
    line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line_p.paragraph_format.space_after = Pt(20)
    lr = line_p.add_run('─' * 28)
    lr.font.size = Pt(10)
    lr.font.color.rgb = LIGHT_BLUE

    # 標題（動態字型大小）
    title = meta.get('title', meta.get('topic', '研究報告'))
    title_len = len(title)
    if title_len <= 20:
        title_size = 28
    elif title_len <= 30:
        title_size = 22
    elif title_len <= 40:
        title_size = 18
    else:
        title_size = 15

    cover_para(title, title_size, bold=True, color=WHITE,
               space_before=4, space_after=10)

    # 副標
    subtitle = meta.get('subtitle', '')
    if subtitle:
        cover_para(subtitle, 11, color=LIGHT_BLUE, space_after=26)

    # 裝飾線
    line_p2 = cell.add_paragraph()
    line_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line_p2.paragraph_format.space_after = Pt(20)
    lr2 = line_p2.add_run('─' * 28)
    lr2.font.size = Pt(10)
    lr2.font.color.rgb = RGBColor(0x00, 0x50, 0x90)

    # 日期
    date_str = meta.get('date', datetime.now().strftime('%Y年%-m月'))
    cover_para(date_str, 11, color=LIGHT_GREY, space_before=6, space_after=4)
    cover_para('本報告由 Dolphin.Ai 輔助生成', 8, color=RGBColor(0x77, 0x88, 0xAA),
               space_after=12)

    doc.add_page_break()


# ── Executive Summary ─────────────────────────────

def build_executive_summary(doc, summary):
    # 區塊標題
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(14)
    add_run(p, '主管摘要', bold=True, size=20, color=NEXI_BLUE)

    # 核心結論（帶淡藍背景框）
    conclusion = summary.get('core_conclusion', '')
    if conclusion:
        table = doc.add_table(rows=1, cols=1)
        remove_table_borders(table)
        table.autofit = False
        table.columns[0].width = Cm(15)
        cell = table.cell(0, 0)
        set_cell_bg(cell, RGBColor(0xE8, 0xF4, 0xFF))
        set_cell_margins(cell, top=200, bottom=200, left=300, right=300)
        cell.paragraphs[0].clear()
        cp = cell.add_paragraph()
        cp.paragraph_format.space_after = Pt(0)
        add_run(cp, conclusion, size=11, color=CHARCOAL)
        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 關鍵發現
    findings = summary.get('key_findings', [])
    if findings:
        lp = doc.add_paragraph()
        add_run(lp, '關鍵發現', bold=True, size=12, color=NEXI_BLUE)
        lp.paragraph_format.space_after = Pt(6)
        for finding in findings:
            fp = doc.add_paragraph()
            fp.paragraph_format.left_indent  = Cm(0.5)
            fp.paragraph_format.space_after  = Pt(4)
            add_run(fp, '▸ ', bold=True, color=TEAL, size=11)
            add_run(fp, finding, size=11, color=CHARCOAL)

    # 建議行動
    recs = summary.get('recommendations', [])
    if recs:
        doc.add_paragraph().paragraph_format.space_after = Pt(4)
        rp = doc.add_paragraph()
        add_run(rp, '建議行動', bold=True, size=12, color=NEXI_BLUE)
        rp.paragraph_format.space_after = Pt(6)
        for i, rec in enumerate(recs, 1):
            rpp = doc.add_paragraph()
            rpp.paragraph_format.left_indent = Cm(0.5)
            rpp.paragraph_format.space_after = Pt(4)
            add_run(rpp, f'{i}. ', bold=True, color=AMBER, size=11)
            add_run(rpp, rec, size=11, color=CHARCOAL)

    doc.add_paragraph()


# ── Section ───────────────────────────────────────

def build_section(doc, section):
    # 章節標題列（帶左側色條用表格模擬）
    title_table = doc.add_table(rows=1, cols=2)
    remove_table_borders(title_table)
    title_table.autofit = False
    title_table.columns[0].width = Cm(0.4)
    title_table.columns[1].width = Cm(14.6)

    bar_cell = title_table.cell(0, 0)
    set_cell_bg(bar_cell, NEXI_BLUE)
    set_row_height(title_table.rows[0], 560, exact=False)

    title_cell = title_table.cell(0, 1)
    set_cell_margins(title_cell, top=80, bottom=80, left=200, right=0)
    title_cell.paragraphs[0].clear()
    tp = title_cell.add_paragraph()
    tp.paragraph_format.space_before = Pt(2)
    tp.paragraph_format.space_after  = Pt(2)
    add_run(tp, section.get('title', ''), bold=True, size=15, color=NEXI_BLUE)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # 正文：先清理 markdown 語法，再切段
    content = _strip_md_block(section.get('content', ''))

    raw_blocks = [b.strip() for b in content.split('\n\n') if b.strip()]
    if len(raw_blocks) <= 1:
        # Claude 可能只用單換行，或根本沒換行 — 嘗試單換行切分
        raw_blocks = [b.strip() for b in content.split('\n') if b.strip()]

    # 自動斷段：若任何段落超過 350 字，在句號處切開
    paragraphs = []
    for block in raw_blocks:
        if len(block) <= 350:
            paragraphs.append(block)
        else:
            sentences = _re.split(r'(?<=[。！？])\s*', block)
            current = ''
            for sent in sentences:
                if len(current) + len(sent) > 300 and current:
                    paragraphs.append(current.strip())
                    current = sent
                else:
                    current += sent
            if current.strip():
                paragraphs.append(current.strip())

    for para_text in paragraphs:
        _render_rich_paragraph(doc, para_text)

    # 資料表格（由 Reporter 生成的 [TABLE_JSON] 解析而來）
    tables = section.get('tables', [])
    for tbl_data in tables:
        build_data_table(doc, tbl_data)

    # 關鍵數據框
    key_data = [d for d in section.get('key_data', []) if d and d.strip()]
    if key_data:
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
        dt = doc.add_table(rows=1, cols=1)
        remove_table_borders(dt)
        dt.autofit = False
        dt.columns[0].width = Cm(14)
        dc = dt.cell(0, 0)
        set_cell_bg(dc, DATA_BG)
        set_cell_margins(dc, top=150, bottom=150, left=250, right=250)
        dc.paragraphs[0].clear()

        label_p = dc.add_paragraph()
        label_p.paragraph_format.space_after = Pt(4)
        add_run(label_p, '▌ 關鍵數據', bold=True, color=NEXI_BLUE, size=10)

        for d in key_data:
            dp = dc.add_paragraph()
            dp.paragraph_format.space_after = Pt(3)
            add_run(dp, '· ', color=TEAL, size=10)
            add_run(dp, d, color=CHARCOAL, size=10)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ── Risk Section ──────────────────────────────────

def build_risk_section(doc, risk_data):
    doc.add_page_break()
    p = doc.add_paragraph()
    add_run(p, '研究限制與風險提示', bold=True, size=16, color=NEXI_BLUE)
    p.paragraph_format.space_after = Pt(12)

    labels = {
        'information_gaps': '資訊缺口',
        'key_assumptions':  '關鍵假設前提',
        'counter_arguments':'反面觀點',
    }
    for key, label in labels.items():
        items = risk_data.get(key, [])
        if not items:
            continue
        lp = doc.add_paragraph()
        add_run(lp, label, bold=True, size=12, color=CHARCOAL)
        lp.paragraph_format.space_after = Pt(4)
        for item in items:
            ip = doc.add_paragraph()
            ip.paragraph_format.left_indent = Cm(0.5)
            ip.paragraph_format.space_after = Pt(4)
            add_run(ip, '▸ ', bold=True, color=LIGHT_BLUE, size=11)
            add_run(ip, item, size=11, color=CHARCOAL)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ── Sources ───────────────────────────────────────

def build_sources_appendix(doc, sources):
    doc.add_page_break()

    # 標題列（帶底線）
    p = doc.add_paragraph()
    add_run(p, '附錄：參考資料來源', bold=True, size=16, color=NEXI_BLUE)
    p.paragraph_format.space_after = Pt(4)

    # 來源數量小字
    meta_p = doc.add_paragraph()
    add_run(meta_p, f'共 {len(sources)} 筆來源　·　由 Dolphin.Ai 自動蒐集', size=9, color=LIGHT_GREY)
    meta_p.paragraph_format.space_after = Pt(14)

    for i, s in enumerate(sources, 1):
        # 標題行
        np = doc.add_paragraph()
        np.paragraph_format.space_after = Pt(1)
        add_run(np, f'[{i}]  ', bold=True, size=10, color=LIGHT_BLUE)
        add_run(np, s.get('title', '未知來源'), bold=True, size=10, color=CHARCOAL)

        # URL 行
        url = s.get('url', '')
        if url:
            up = doc.add_paragraph()
            up.paragraph_format.left_indent = Cm(0.7)
            up.paragraph_format.space_after = Pt(1)
            add_run(up, url, size=9, color=LIGHT_BLUE)

        # 擷取日期（JS 輸出為 fetched_at，相容 accessed）
        accessed = s.get('accessed') or s.get('fetched_at', '')
        if accessed:
            ap = doc.add_paragraph()
            ap.paragraph_format.left_indent = Cm(0.7)
            ap.paragraph_format.space_after = Pt(8)
            add_run(ap, f'擷取日期：{accessed}', size=8, color=LIGHT_GREY)
        else:
            doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ── Header / Footer ───────────────────────────────

def add_page_number(paragraph):
    """在段落中插入 Word 自動頁碼 field code"""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run._r.append(instrText)

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)


def add_header_footer(doc, title):
    short_title = title if len(title) <= 30 else title[:28] + '…'
    for sec in doc.sections:
        # 頁首
        header = sec.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.clear()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_run(hp, f'Dolphin.Ai  |  {short_title}', size=8, color=LIGHT_GREY)

        # 頁尾 — 含自動頁碼
        footer = sec.footer
        footer.is_linked_to_previous = False
        if not footer.paragraphs:
            footer.add_paragraph()
        fp = footer.paragraphs[0]
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(fp, '本報告由 Dolphin.Ai 輔助生成  ·  僅供參考  ·  第 ', size=8, color=LIGHT_GREY)
        add_page_number(fp)
        add_run(fp, ' 頁', size=8, color=LIGHT_GREY)


# ── Main ──────────────────────────────────────────

def generate_report(report_content: dict, output_path: str):
    doc = Document()

    # 頁面設定
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(3.0)
        sec.right_margin  = Cm(2.5)

    meta  = report_content.get('meta', {})
    title = meta.get('title', meta.get('topic', '研究報告'))

    # 封面
    build_cover_page(doc, meta)

    # 目錄頁
    sections_list = report_content.get('sections', [])
    summary = report_content.get('executive_summary', {})
    build_toc(doc, sections_list, summary_exists=bool(summary))

    # 主管摘要
    if summary:
        build_executive_summary(doc, summary)

    # 章節分隔頁（淡色標題）
    for section in report_content.get('sections', []):
        build_section(doc, section)

    # 風險章節
    risk = report_content.get('risk_and_limitations', {})
    if risk and any(risk.values()):
        build_risk_section(doc, risk)

    # 來源附錄
    sources = report_content.get('sources', [])
    if sources:
        build_sources_appendix(doc, sources)

    # 頁首頁尾
    add_header_footer(doc, title)

    doc.save(output_path)
    print(f'[DOCX] ✓ 報告已生成：{output_path}')
    return output_path


def generate_summary_card(report_content: dict, output_path: str):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Cm(2)
        sec.left_margin = sec.right_margin = Cm(2.5)

    meta    = report_content.get('meta', {})
    summary = report_content.get('executive_summary', {})

    # 頂部色帶
    table = doc.add_table(rows=1, cols=1)
    remove_table_borders(table)
    table.autofit = False
    table.columns[0].width = Cm(16)
    hcell = table.cell(0, 0)
    set_cell_bg(hcell, NEXI_BLUE)
    set_cell_margins(hcell, top=300, bottom=300, left=400, right=400)
    set_row_height(table.rows[0], 1600, exact=False)
    hcell.paragraphs[0].clear()

    bp = hcell.add_paragraph()
    bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(bp, 'Dolphin.Ai', bold=True, size=11, color=TEAL)

    title = meta.get('title', meta.get('topic', '研究摘要'))
    title_size = 18 if len(title) <= 25 else (14 if len(title) <= 35 else 12)
    tp = hcell.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_before = Pt(4)
    add_run(tp, title, bold=True, size=title_size, color=WHITE)

    sub = meta.get('subtitle', '')
    if sub:
        sp = hcell.add_paragraph()
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(sp, sub, size=9, color=LIGHT_BLUE)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 核心結論
    conclusion = summary.get('core_conclusion', '')
    if conclusion:
        cp = doc.add_paragraph()
        add_run(cp, '核心結論', bold=True, size=13, color=NEXI_BLUE)
        cp.paragraph_format.space_after = Pt(6)
        body_p = doc.add_paragraph()
        add_run(body_p, conclusion, size=11, color=CHARCOAL)
        body_p.paragraph_format.space_after = Pt(12)

    # 關鍵發現
    findings = summary.get('key_findings', [])
    if findings:
        fp = doc.add_paragraph()
        add_run(fp, '關鍵發現', bold=True, size=13, color=NEXI_BLUE)
        fp.paragraph_format.space_after = Pt(6)
        for i, f in enumerate(findings[:5], 1):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.4)
            p.paragraph_format.space_after = Pt(5)
            add_run(p, f'{i}.  ', bold=True, size=11, color=AMBER)
            add_run(p, f, size=11, color=CHARCOAL)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 頁尾
    efp = doc.add_paragraph()
    efp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(efp, f'Dolphin.Ai  ·  {meta.get("date", "")}  ·  AI 輔助生成', size=8, color=LIGHT_GREY)

    doc.save(output_path)
    print(f'[DOCX] ✓ 摘要卡已生成：{output_path}')
    return output_path


if __name__ == '__main__':
    parser = argparse.ArgumentParser()
    parser.add_argument('--input',      required=True)
    parser.add_argument('--output-dir', default='./output')
    args = parser.parse_args()

    with open(args.input, 'r', encoding='utf-8') as f:
        content = json.load(f)

    os.makedirs(args.output_dir, exist_ok=True)
    topic    = content.get('meta', {}).get('topic', 'report')
    safe     = topic[:20].replace('/', '_').replace(' ', '_')
    date_str = datetime.now().strftime('%Y%m%d')

    generate_report(content,       os.path.join(args.output_dir, f'{safe}_完整報告_{date_str}.docx'))
    generate_summary_card(content, os.path.join(args.output_dir, f'{safe}_摘要卡_{date_str}.docx'))
